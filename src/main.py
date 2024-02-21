# from tkinter import Tk
import pandas as pd
from docx import Document
from datetime import datetime

excel_data = pd.read_excel("src/template/EmployeeList.xlsx",dtype={"ID":str})
date_columns = ["Date","From", "To","DOB"]  # List your date columns

for col in date_columns:
    excel_data[col] = pd.to_datetime(excel_data[col])

# excel_data.fillna(".......", inplace=True)

# Check xem mot cot co phai dang ngay khong
def is_date_column(col):
    return pd.api.types.is_datetime64_any_dtype(col)

# Check xem mot cot co phai la so khong
def is_num_colum(col):
    return pd.api.types.is_numeric_dtype(col)

# Format ngay voi dang ddmmyy
def format_date(date):
    if pd.notnull(date):
        return date.strftime("%d/%m/%Y")
    return ""

# Ngan cach phan nghin cho so
def format_number(num):
    if pd.notnull(num):
        return f"{num:,}".replace(",",".")
    return ""

# The hien so 0 truoc so co 1 chu so
def adjust_number(num):
    if pd.notnull(num) and num < 10 and num > 0:
        return "0"+str(int(num))
    elif pd.notnull(num):
        return str(int(num))

def process_date(date):
    parsed_date = pd.to_datetime(date)
    return {
        "Day":adjust_number(parsed_date.day),
        "Month":adjust_number(parsed_date.month),
        "Year":adjust_number(parsed_date.year)
            }
# Loop tat ca cac dong trong excel
for index,row in excel_data.iterrows():
    if "Date" in row.keys():
        date_data = process_date(row["Date"])
        for key, value in date_data.items():
            row[key] = value
    formatted_row = row.copy()

    document = Document('src/template/Contract.docx')
    document1 = Document('src/template/Contract.docx')

    if row["ContractType"] == "Không xác định thời hạn":
        document = Document('src/template/Contract-LongTerm.docx')
        document1 = Document('src/template/Contract-LongTerm.docx')
    elif row["ContractType"] == "Cộng tác viên":
        document = Document('src/template/Contract-Partner.docx')

    for col in excel_data.columns:
        if is_date_column(excel_data[col]):
            formatted_row[col] = format_date(row[col])
        if is_num_colum(excel_data[col]):
            formatted_row[col] = adjust_number(row[col])
        if excel_data[col] is None or excel_data[col].empty:
            formatted_row[col] = "......"
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            # print(run.text)
            # print("Key: "+key)
            for key in row.keys():
                placeholder = f'{{{key}}}'
                if placeholder in run.text:
                    # print(f"Attempting to replace: {placeholder} with data: {formatted_row[key]}")
                # print(formatted_row[key])
                    if pd.isna(formatted_row[key]) and key=="Add":
                        print(formatted_row["Employee"]+" hợp đồng thường bị trống thông tin excel")
                        # run.text=""
                        run.text = run.text.replace(placeholder,str(""))
                    else:
                        run.text = run.text.replace(placeholder,str(formatted_row[key]))
    document.save(f'HD/HopDong_{formatted_row["No"]}_{row["Employee"]}.docx')

    if row["ContractType"] == "Cộng tác viên":
        continue   
    for paragraph in document1.paragraphs:
        for run in paragraph.runs:
            # print(run.text)
            # print("Key: "+key)
            for key in row.keys():
                placeholder = f'{{{key}}}'
                if placeholder in run.text:
                    if key == "Add" or placeholder=="Add":
                        # run.text=None
                        run.text = run.text.replace(placeholder,"")
                        # paragraph.clear()
                    if key == "Salary":
                        run.text = run.text.replace(placeholder,"3.860.000")
                    elif key == "SalaryText":
                        run.text = run.text.replace(placeholder,"Ba triệu tám trăm sáu mươi nghìn đồng chẵn")
                    else:
                        run.text = run.text.replace(placeholder,str(formatted_row[key]))

    document1.save(f'BaoHiem/HopDong_{formatted_row["No"]}_{row["Employee"]}_BH.docx')